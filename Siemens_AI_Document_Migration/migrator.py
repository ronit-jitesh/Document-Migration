"""
Siemens Document Migration Engine
----------------------------------
Reads unstructured SOP text and:
  1. Extracts structured data via LLM (OpenAI + LangChain)
  2. Generates a clean, standardized Word document (.docx)
"""

import os
import re
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from pypdf import PdfReader

from schema import SiemensSOP

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
load_dotenv()

OUTPUT_DIR = Path(__file__).parent / "output_docs"
OUTPUT_DIR.mkdir(exist_ok=True)

EXTRACTION_PROMPT = """\
You are a senior Quality Assurance Document Specialist at Siemens Digital Industries.

Your task is to read the following messy, unstructured Standard Operating Procedure (SOP) \
and extract ALL relevant information into a clean, standardized JSON format.

Rules:
1. **title**: Create a professional, standardized title for this procedure.
2. **document_id**: Generate a plausible document ID in the format "SOP-YYYY-NNN".
3. **version**: Assign version "2.0" (this is the first standardized migration).
4. **department**: Infer the responsible department from context.
5. **safety_warnings**: Extract EVERY safety warning, caution, danger notice, PPE requirement, \
   and EHS reference. Do NOT miss any. Rewrite them clearly and professionally.
6. **equipment**: List ALL tools, instruments, and equipment mentioned. Standardize names.
7. **steps**: Extract the procedural steps in logical order. Clean up language, \
   remove informal tone, and write each step as a clear, actionable instruction.
8. **confidence_score**: Rate 1-10 how confident you are that you captured ALL information \
   from the source document. Be honest — if the source was very messy, a score of 7-8 is fine.

Source Document:
---
{document_text}
---

Extract the data now. Return ONLY the structured JSON."""


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------
def read_file(file_path: str) -> str:
    """Read text from a .txt or .pdf file."""
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        return path.read_text(encoding="utf-8")


def extract_sop(text: str, model: str = "gpt-4o-mini") -> SiemensSOP:
    """
    Send raw document text to the LLM and receive a validated SiemensSOP object.
    Uses LangChain's structured output to guarantee Pydantic compliance.
    """
    llm = ChatOpenAI(model=model, temperature=0)
    structured_llm = llm.with_structured_output(SiemensSOP)

    prompt = ChatPromptTemplate.from_messages(
        [("system", EXTRACTION_PROMPT)]
    )

    chain = prompt | structured_llm
    result = chain.invoke({"document_text": text})
    return result


# ---------------------------------------------------------------------------
# Word Document Generation
# ---------------------------------------------------------------------------
def _add_heading(doc: Document, text: str, level: int = 0):
    """Add a formatted heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0x80, 0x80)  # Siemens teal


def _add_section_header(doc: Document, text: str):
    """Add a bold section header with a horizontal rule."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark navy


def generate_word_doc(data: SiemensSOP) -> str:
    """
    Generate a professional Siemens-branded Word document from a SiemensSOP object.
    Returns the absolute path to the saved .docx file.
    """
    doc = Document()

    # --- Styles setup ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Header ---
    _add_heading(doc, "Siemens Digital Industries", level=0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Standard Operating Procedure")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.bold = True

    doc.add_paragraph()  # spacer

    # --- Metadata Table ---
    table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    metadata = [
        ("Document Title", data.title),
        ("Document ID", data.document_id),
        ("Version", data.version),
        ("Department", data.department),
    ]
    for i, (label, value) in enumerate(metadata):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # --- Safety Warnings ---
    _add_section_header(doc, "⚠  Safety Warnings & Precautions")
    for warning in data.safety_warnings:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(warning)
        run.font.color.rgb = RGBColor(0xCC, 0, 0)  # Red
        run.bold = True
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # --- Equipment ---
    _add_section_header(doc, "🔧  Required Equipment")
    for item in data.equipment:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # --- Procedural Steps ---
    _add_section_header(doc, "📋  Procedure Steps")
    for i, step in enumerate(data.steps, 1):
        para = doc.add_paragraph()
        num_run = para.add_run(f"Step {i}:  ")
        num_run.bold = True
        num_run.font.size = Pt(10)
        num_run.font.color.rgb = RGBColor(0, 51, 102)
        step_run = para.add_run(step)
        step_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # --- Footer ---
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "This document was auto-generated by the Siemens AI Document Migration System. "
        "Please review and approve before use."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # --- Save ---
    safe_title = re.sub(r"[^\w\s-]", "", data.title).strip().replace(" ", "_")[:50]
    filename = f"{safe_title}_v{data.version}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))

    return str(filepath)


# ---------------------------------------------------------------------------
# CLI entry point (for quick testing)
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python migrator.py <path_to_input_file>")
        sys.exit(1)

    input_file = sys.argv[1]
    print(f"📄 Reading: {input_file}")
    raw_text = read_file(input_file)
    print(f"📝 Extracted {len(raw_text)} characters")

    print("🤖 Sending to LLM for structured extraction...")
    sop = extract_sop(raw_text)
    print(f"✅ Extracted SOP: {sop.title}")
    print(f"   Confidence: {sop.confidence_score}/10")
    print(f"   Safety Warnings: {len(sop.safety_warnings)}")
    print(f"   Steps: {len(sop.steps)}")

    print("📝 Generating Word document...")
    output_path = generate_word_doc(sop)
    print(f"✅ Saved to: {output_path}")
